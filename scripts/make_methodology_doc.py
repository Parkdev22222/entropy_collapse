#!/usr/bin/env python3
# Copyright 2026 STEER-F authors
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
"""Generate docs/STEER-F_methodology.docx — the paper's Method section.

Written as a script rather than a hand-made binary so the document can be
regenerated after the code changes, and so its claims can be traced: every
formula and default here is taken from the implementation, not from the plan
document, which `docs/steer_code_map.md` records as diverging from what STEER
actually does.

    python scripts/make_methodology_doc.py [--out docs/STEER-F_methodology.docx]
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

BODY_FONT = "Times New Roman"
MONO_FONT = "Consolas"


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    code = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = MONO_FONT
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Pt(24)
    code.paragraph_format.space_before = Pt(6)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing = 1.0

    eq = doc.styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
    eq.font.name = MONO_FONT
    eq.font.size = Pt(10)
    eq.paragraph_format.left_indent = Pt(24)
    eq.paragraph_format.space_before = Pt(8)
    eq.paragraph_format.space_after = Pt(8)

    cap = doc.styles.add_style("Caption2", WD_STYLE_TYPE.PARAGRAPH)
    cap.font.name = BODY_FONT
    cap.font.size = Pt(9)
    cap.font.italic = True
    cap.paragraph_format.space_after = Pt(10)

    todo = doc.styles.add_style("Placeholder", WD_STYLE_TYPE.PARAGRAPH)
    todo.font.name = BODY_FONT
    todo.font.size = Pt(10.5)
    todo.font.italic = True
    todo.font.color.rgb = RGBColor(0x99, 0x33, 0x00)
    todo.paragraph_format.left_indent = Pt(18)
    todo.paragraph_format.space_after = Pt(10)


def para(doc, text, style=None):
    p = doc.add_paragraph(text, style=style)
    return p


def code(doc, text):
    for line in text.strip("\n").split("\n"):
        doc.add_paragraph(line, style="CodeBlock")


def equation(doc, text):
    for line in text.strip("\n").split("\n"):
        p = doc.add_paragraph(line, style="Equation")
    return p


def caption(doc, text):
    doc.add_paragraph(text, style="Caption2")


def placeholder(doc, text):
    doc.add_paragraph(text, style="Placeholder")


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = BODY_FONT
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(v))
            run.font.size = Pt(9)
            run.font.name = MONO_FONT if i == 0 and len(str(v)) < 24 else BODY_FONT
    doc.add_paragraph()
    return t


def build(out_path):
    doc = Document()
    setup_styles(doc)

    # ---------------------------------------------------------------- title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("STEER-F: Correcting Myopic Entropy Control in RLVR\n"
                      "with Multi-Token Future-Entropy Forecasting")
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = BODY_FONT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Method (draft — experimental results pending)")
    rs.italic = True
    rs.font.size = Pt(11)

    doc.add_paragraph()

    # ---------------------------------------------------------------- abstract
    doc.add_heading("Abstract", level=1)
    para(doc,
         "Entropy-aware reweighting stabilises reinforcement learning with verifiable "
         "rewards (RLVR) by scaling each token's learning signal according to how the "
         "update is predicted to change that token's entropy. We observe that this "
         "estimate is myopic: it captures only the local term of the trajectory-entropy "
         "decomposition and is blind to the visitation term, under which a policy that "
         "concentrates probability at a branch point deletes every future state along the "
         "abandoned branch without any local entropy change large enough to notice. "
         "STEER-F supplies the missing term. Multi-token prediction (MTP) heads forecast "
         "the entropy that lies ahead of each decision; the forecast is converted into a "
         "branch score by subtracting a prefix-matched sibling baseline; and the score "
         "extends the reweighting metric. The construction is exactly reducible: at "
         "λ = 0 the method is bit-identical to the stock entropy-control baseline, "
         "which we enforce by test against a verbatim copy of the upstream function rather "
         "than by inspection, so that every reported difference is attributable to the "
         "future-entropy term alone.")
    placeholder(doc, "[Abstract to be completed with headline numbers once Phase 2 finishes.]")

    # ---------------------------------------------------------------- 1
    doc.add_heading("1  Preliminaries", level=1)

    doc.add_heading("1.1  RLVR with group-relative advantages", level=2)
    para(doc,
         "We train a policy πθ on verifiable mathematical reasoning with GRPO. For "
         "each prompt, a group of n rollouts is sampled; the reward is a binary "
         "verifier outcome; and the advantage of token t in rollout i is the "
         "group-standardised return. The policy objective is the usual clipped "
         "surrogate with asymmetric clipping bounds, no KL penalty to a reference "
         "policy, and no entropy bonus, so that entropy dynamics are governed entirely "
         "by the reweighting under study.")

    doc.add_heading("1.2  The local entropy-change estimate", level=2)
    para(doc,
         "The baseline method reweights each token by a first-order prediction of the "
         "entropy change that one gradient step induces at that token. Writing π for "
         "the current probability of the sampled token, π_old for its probability "
         "under the behaviour policy, A for the advantage and H for the token's policy "
         "entropy, the estimate is")
    equation(doc,
             "x       = clamp(exp(log π), 1e-8, 1 - 1e-8)\n"
             "f       = x · (1 - x) · (log x + H)\n"
             "Ω_local = - ( A / clamp(exp(log π_old), 1e-8, 1) ) · f")
    para(doc,
         "with the sign convention that Ω > 0 means the update is predicted to raise "
         "entropy at this position. Ω is then mapped to a bounded token weight "
         "w ∈ [w_min, w_max] and multiplies the per-token policy-gradient loss.")

    doc.add_heading("1.3  What the mapping actually is", level=2)
    para(doc,
         "This subsection records a discrepancy that determines several design choices "
         "downstream, and we state it explicitly because the natural reading of the "
         "baseline's description does not match its implementation.")
    para(doc,
         "There is no entropy band [ΔH_low, ΔH_high] and no discrete multiplier "
         "α ∈ {γ, 1, 1/γ}. The implementation rescales Ω onto "
         "[w_min, w_max] per micro-batch, by one of two mappings:")
    table(doc,
          ["Mapping", "Form", "Invariant under"],
          [["linear=True",
            "min–max rescale of Ω onto [w_min, w_max]",
            "any affine a·x + b with a > 0"],
           ["linear=False",
            "w = exp(±k·Ω), k = -log w_min / max(|Ω|_max, 0.02)",
            "positive scaling a·x only (exp is not shift-invariant)"]])
    para(doc,
         "Additionally, mode=\"symmetric\" applies abs(·) to the metric before the "
         "mapping, so the point zero carries meaning and any recentring changes which "
         "tokens count as large entropy changes.")
    para(doc,
         "Two consequences follow. First, the min–max statistics are computed per "
         "micro-batch, so the micro-batch size is part of the method's configuration "
         "rather than a memory knob; it must be pinned across every arm being compared. "
         "Second, the only normalisation that leaves all four (mode, mapping) "
         "combinations untouched is positive scaling without recentring — which is "
         "why §2.4 uses RMS scaling rather than a z-score.")

    # ---------------------------------------------------------------- 2
    doc.add_heading("2  Method", level=1)

    doc.add_heading("2.1  Motivation: the missing visitation term", level=2)
    para(doc,
         "Trajectory entropy decomposes into a local term — the entropy of the "
         "next-token distribution at each visited state — and a visitation term, "
         "the entropy induced by which states are reached at all. Ω_local sees only "
         "the first. At a branch point, concentrating probability on one continuation "
         "may change the local next-token entropy very little while removing an entire "
         "subtree of future states from the policy's support. STEER-F adds an estimate "
         "of that second effect.")

    doc.add_heading("2.2  Forecasting future entropy with MTP heads", level=2)
    para(doc,
         "We attach K lightweight heads to the frozen policy body. Head k reads the "
         "policy's final hidden state h_t and predicts the distribution of the token at "
         "offset t + k + 1. Each head is a bottleneck MLP with a residual connection,")
    equation(doc,
             "head_k(h) = h + W2_k · SiLU( W1_k · h ),      W1_k ∈ R^{d×H}, W2_k ∈ R^{H×d}\n"
             "logits_k  = lm_head( head_k(h) )")
    para(doc,
         "The unembedding is shared with the policy, so the heads add only "
         "2·K·H·d parameters and inherit a calibrated output geometry. The "
         "final projection of every head is zero-initialised, making head k an "
         "identity-initialised copy of the +1 predictor at the start of training.")
    para(doc,
         "Entropy is computed without ever materialising the [K, B, T, V] logit stack: "
         "positions are chunked through the unembedding and reduced to a scalar entropy "
         "in place, so peak additional memory is chunk_size × V and is independent "
         "of sequence length.")

    doc.add_heading("2.3  From forecasts to a branch score", level=2)
    para(doc, "The discounted entropy-to-go at a state s is")
    equation(doc,
             "H_togo(s) = Σ_{k=1..κ}  γ_H^k · H( p_MTP( y_{t+k} | s ) )")
    para(doc,
         "The heads are trained with K ≥ κ and are independent, so the horizon "
         "κ is selected post hoc without retraining.")
    para(doc,
         "Index alignment is load-bearing. The hidden state at position t is the one "
         "produced after consuming y_t, so a forecast read at that position is "
         "conditioned on the sampled token: it is H_togo(s_t ⊕ y_t), not "
         "H_togo(s_t). The branch score subtracts a baseline that conditions on the "
         "state only:")
    equation(doc, "A_H(s_t, y_t) = H_togo(s_t ⊕ y_t) − H̄_togo(s_t)")
    para(doc, "We consider two baselines.")
    table(doc,
          ["Baseline", "Definition"],
          [["sibling (default)",
            "Mean H_togo over rollouts in the same prompt group whose responses agree "
            "on every token strictly before t. The strict inequality is what makes the "
            "difference a branch score: siblings must share the prefix but are free to "
            "differ at t itself."],
           ["group",
            "Mean H_togo over the whole prompt group at position t, ignoring prefix "
            "agreement."]])
    para(doc,
         "A property of the sibling baseline is worth stating because it governs how "
         "A_H must be summarised. Once a rollout has become unique within its group it "
         "is its own only sibling, so its baseline equals its own value and A_H is "
         "exactly zero. Past the first divergence this holds at nearly every position: "
         "with eight rollouts sharing a short prefix, A_H is exactly zero at 93.8% of "
         "positions at T = 96, 98.8% at T = 512, and 99.4% at T = 1024. A_H is "
         "therefore a sparse signal supported on the sibling-rich region, and any "
         "statistic defined as “the top decile of A_H” must select by rank "
         "rather than by thresholding at the k-th largest value, since a value "
         "threshold admits the entire tie block at zero.")
    caption(doc,
            "Table 1. Zero mass of A_H under the sibling baseline, measured on eight "
            "siblings with a shared five-token prefix.")

    doc.add_heading("2.4  The extended metric", level=2)
    para(doc,
         "A gradient step on the clipped surrogate moves the sampled token's logit by "
         "η·w; the softmax Jacobian turns this into a first-order shift of its "
         "log-probability:")
    equation(doc, "Δlogπ̂ = η · w · (1 − π)")
    para(doc,
         "Multiplying the shift by the clipped branch score gives the visitation "
         "contribution, and the extended metric is the normalised sum of the two terms:")
    equation(doc,
             "Ω̃ = norm(Ω_local) + λ · norm( Δlogπ̂ · clip(A_H, −c, c) )")
    para(doc,
         "The sign reading is consistent with Ω_local. When Δlogπ̂ > 0 "
         "and A_H > 0 the update concentrates mass on a branch whose future is more "
         "diverse than its siblings', so trajectory entropy rises; when A_H < 0 mass "
         "moves onto a comparative dead end and trajectory entropy falls, even though "
         "the local term may register nothing.")
    para(doc,
         "norm(·) divides by the masked RMS. As established in §1.3, positive "
         "scaling without recentring is the only normalisation compatible with all four "
         "(mode, mapping) combinations of the baseline; a z-score would break the "
         "symmetric mode, where abs(·) makes zero a meaningful pivot. A degenerate "
         "input returns zeros rather than dividing by a near-zero scale, which is safe "
         "because an all-equal metric already maps every token to the same weight.")
    para(doc,
         "Boundedness is provided by clip(A_H, −c, c) together with the clamp onto "
         "[w_min, w_max] that the baseline mapping already applies; no discrete "
         "multiplier is involved, consistent with §1.3.")

    doc.add_heading("2.5  Exact reducibility at λ = 0", level=2)
    para(doc,
         "At λ = 0 the implementation returns Ω_local itself, before any "
         "normalisation is applied, rather than relying on norm(·) happening to be a "
         "no-op. The λ = 0 arm therefore traverses the same code path as λ > 0 "
         "while producing bit-identical token weights to the unmodified upstream "
         "function. This is verified by test against a verbatim copy of the upstream "
         "implementation across all four (mode, mapping) combinations, and the test "
         "suite additionally demonstrates the failure mode that a recentring "
         "normalisation would introduce. The consequence for the experimental design is "
         "that the STEER-F effect is a controlled within-codebase comparison rather than "
         "a comparison across implementations.")

    doc.add_heading("2.6  Head calibration", level=2)
    para(doc,
         "Distant heads systematically over-estimate entropy: their predictive "
         "distribution is blurred by everything that may intervene, so H drifts toward "
         "log V as k grows. Uncorrected, this makes the far future look uniformly "
         "diverse and destroys the discriminative power of A_H. We fit a per-head "
         "correction against Monte-Carlo measured entropy,")
    equation(doc,
             "H_cal,k = scale_k · H( softmax( logits_k / τ_k ) ) + bias_k")
    para(doc,
         "by least squares, with τ folded into the entropy computation itself and "
         "(scale, bias) as the residual affine fit. Heads whose forecast has no variance "
         "receive scale = 1, bias = 0 instead of a division by zero. Calibration is "
         "re-estimated per model family; all other hyper-parameters are transferred "
         "unchanged, since transferability is itself a claim under test.")

    # ---------------------------------------------------------------- 3
    doc.add_heading("3  Validation protocol for the forecast", level=1)
    para(doc,
         "Because the entire method rests on A_H carrying information, we validate the "
         "forecast offline before committing it to RL, and treat a failure here as a "
         "publishable negative result rather than as an obstacle to be tuned away.")

    doc.add_heading("3.1  Criterion 1 — rank correlation with measured entropy", level=2)
    para(doc,
         "For each of a set of prefixes drawn from on-policy trajectories, we sample "
         "continuations and measure the policy's entropy at each offset, then compare "
         "the forecast H_togo against the measured counterpart. Two details matter. "
         "First, the comparison applies the same discount γ_H to the measurement "
         "as to the forecast; scoring a discounted forecast against an undiscounted "
         "measurement would penalise γ_H < 1 for reasons unrelated to forecast "
         "quality. Second, correlation is computed within problem and averaged over "
         "problems by Fisher-z, since between-problem variation in difficulty would "
         "otherwise dominate a pooled coefficient.")

    doc.add_heading("3.2  Criterion 2 — branch recall against chance", level=2)
    para(doc,
         "A position t of rollout i is a true branch point when at least one rollout "
         "sharing the prefix before t chose a different token at t. This quantity is "
         "defined by siblings and therefore does not exist for independently sampled "
         "continuations; it is measured on real rollout groups. We test whether true "
         "branch points fall inside the top decile of A_H more often than chance, using "
         "an exact binomial test whose null is the realised selection rate rather than "
         "the nominal decile — rounding and ties make the achieved fraction differ, "
         "and testing against the nominal value would bias the verdict.")
    para(doc,
         "Two caveats are reported alongside the statistic. When the number of true "
         "branch points is small the test has little power, so a failure indicates "
         "insufficient sampling rather than an uninformative forecast. And because A_H "
         "is supported on the sibling-rich region (§2.3), its support alone is "
         "partially informative about branch locations; we therefore report the same "
         "statistic computed with untrained heads as a control, so that a pass is "
         "attributable to the forecast's ranking rather than to its support.")

    doc.add_heading("3.3  Diagnostics logged during RL", level=2)
    table(doc,
          ["Quantity", "What it establishes"],
          [["actor/entropy", "Overall policy entropy — the collapse being prevented."],
           ["branch_entropy, branch_entropy_gap",
            "Entropy restricted to the top-decile-A_H tokens versus the rest. A rise "
            "here under a flat overall entropy is the signature of the intended "
            "mechanism rather than of a uniform entropy bonus."],
           ["visit_rel_mag",
            "Magnitude of the visitation term relative to the local term; establishes "
            "that λ is doing neither nothing nor everything."],
           ["a_h_clip_frac",
            "Saturation of clip(A_H, ±c). A saturated clip cannot discriminate "
            "branches."],
           ["token weight occupancy",
            "Fraction of weights in each third of [w_min, w_max]; a strong skew "
            "indicates the mapping range needs resetting."]])

    # ---------------------------------------------------------------- 4
    doc.add_heading("4  Implementation", level=1)
    para(doc,
         "The method is implemented as a drop-in replacement for the baseline's token "
         "weighting, delivered as patches against a pinned upstream commit so that the "
         "complete set of modifications remains a readable diff. All logic with "
         "behavioural content lives outside the patched tree and is unit-tested without "
         "the training framework installed; the patches themselves are mechanical "
         "delegations.")
    para(doc,
         "The forecast requires the policy's final hidden states, which the training "
         "framework's existing log-probability pass discards. Rather than restructuring "
         "that pass, we perform one additional forward under no_grad; there is no "
         "backward, so the per-step cost is approximately one extra forward.")

    doc.add_heading("4.1  Default hyper-parameters", level=2)
    table(doc,
          ["Symbol", "Meaning", "Default"],
          [["λ", "weight of the visitation term (0 = baseline)", "swept"],
           ["η", "logit-step scale in Δlogπ̂", "1.0"],
           ["c", "symmetric clip on A_H", "1.0"],
           ["norm", "normalisation of the two terms", "RMS scale"],
           ["κ", "forecast horizon", "selected in §3"],
           ["γ_H", "entropy-to-go discount", "selected in §3"],
           ["baseline", "A_H baseline", "sibling"],
           ["K", "number of MTP heads trained", "8"],
           ["d", "head bottleneck width", "1024"]])
    para(doc,
         "Only the product λ·η matters after normalisation, so η is "
         "held at 1.0 and λ is swept.")

    doc.add_heading("4.2  Scope of the present implementation", level=2)
    para(doc,
         "Two components of the design are specified and unit-tested but not yet "
         "connected to the training loop: the MTP auxiliary loss that would let the "
         "heads co-evolve with the policy, and the drift controller that would decay "
         "λ when the heads and the policy disagree. Both need the training forward "
         "pass to expose hidden states. The runs reported here therefore hold the heads "
         "frozen at their warm-up state, which coincides with one of the planned "
         "ablations; we state this rather than leave it to be inferred, because it "
         "bounds what the results can support.")

    # ---------------------------------------------------------------- 5
    doc.add_heading("5  Experimental setup", level=1)
    placeholder(doc,
                "[To be completed. Will state: models and scales; datasets and "
                "verifier; the pinned micro-batch size and why it is pinned; the "
                "λ grid; seeds per arm and the resulting resolution relative to "
                "the effect size being claimed; the evaluation protocol and the number "
                "of samples used for pass@k.]")

    doc.add_heading("6  Results", level=1)
    placeholder(doc,
                "[To be completed once the runs finish. Will report: reproduction of "
                "the baseline's entropy behaviour; forecast validation outcomes "
                "including the untrained-head control; the λ sweep with per-seed "
                "values and seed means; the branch-entropy mechanism check; transfer "
                "to other model families; and ablations. Negative and failed runs are "
                "to be reported alongside successful ones.]")

    doc.add_heading("7  Limitations", level=1)
    placeholder(doc,
                "[To be completed. Known items to carry forward: heads are frozen "
                "during RL (§4.2); A_H is exactly zero wherever a rollout has no "
                "siblings, so the correction is inactive over most of a long response; "
                "calibration is fit on the warm-up policy and is not refreshed as the "
                "policy moves; and the sibling baseline makes A_H a deviation from a "
                "group mean, so roughly half of the rollouts at any branch point "
                "receive a negative score by construction.]")

    doc.add_paragraph()
    foot = doc.add_paragraph()
    fr = foot.add_run(
        "Draft generated from the implementation by scripts/make_methodology_doc.py. "
        "Formulas and defaults reflect the code as committed, not the design document; "
        "where the two differ, §1.3 records the difference.")
    fr.italic = True
    fr.font.size = Pt(8.5)

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"[docx] wrote {out_path}")


def main(argv=None):
    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument("--out", default="docs/STEER-F_methodology.docx")
    args = p.parse_args(argv)
    build(args.out)


if __name__ == "__main__":
    main()
